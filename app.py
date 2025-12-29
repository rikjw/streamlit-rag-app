# Core dependencies
import streamlit as st
import openai
import numpy as np
import faiss
from pathlib import Path
from io import BytesIO

# Document processing imports
from pypdf import PdfReader
from docx import Document
import pandas as pd
from bs4 import BeautifulSoup

# ============== Configuration ==============

EMBEDDING_MODEL = "text-embedding-3-small"
CHAT_MODEL = "gpt-4o-mini"
CHUNK_SIZE = 500
CHUNK_OVERLAP = 50

# Supported file types
SUPPORTED_TYPES = ["txt", "md", "pdf", "docx", "csv", "xlsx", "html"]


def get_api_key() -> str | None:
    """Get API key from secrets.toml or environment."""
    # Try Streamlit secrets first
    if "OPENAI_API_KEY" in st.secrets:
        return st.secrets["OPENAI_API_KEY"]

    # Fallback to environment variable
    import os

    if os.getenv("OPENAI_API_KEY"):
        return os.getenv("OPENAI_API_KEY")

    return None


# ============== Document Readers ==============


def read_txt(file) -> str:
    """Read plain text or markdown file."""
    return file.read().decode("utf-8")


def read_pdf(file) -> str:
    """Extract text from PDF file."""
    reader = PdfReader(file)
    text_parts = []

    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            text_parts.append(f"[Page {i+1}]\n{page_text}")

    return "\n\n".join(text_parts)


def read_docx(file) -> str:
    """Extract text from Word document."""
    doc = Document(file)
    paragraphs = []

    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)


def read_csv(file) -> str:
    """Convert CSV to readable text."""
    df = pd.read_csv(file)

    # Create a text representation
    text_parts = [f"CSV Data with {len(df)} rows and {len(df.columns)} columns."]
    text_parts.append(f"Columns: {', '.join(df.columns.tolist())}")
    text_parts.append("\nData:\n")
    text_parts.append(df.to_markdown(index=False))

    return "\n".join(text_parts)


def read_excel(file) -> str:
    """Convert Excel to readable text."""
    xlsx = pd.ExcelFile(file)
    text_parts = []

    for sheet_name in xlsx.sheet_names:
        df = pd.read_excel(xlsx, sheet_name=sheet_name)
        text_parts.append(f"\n### Sheet: {sheet_name}")
        text_parts.append(f"Rows: {len(df)}, Columns: {len(df.columns)}")
        text_parts.append(df.to_markdown(index=False))

    return "\n\n".join(text_parts)


def read_html(file) -> str:
    """Extract text from HTML file."""
    content = file.read()
    soup = BeautifulSoup(content, "html.parser")

    # Remove script and style elements
    for element in soup(["script", "style", "nav", "footer", "header"]):
        element.decompose()

    # Get text with some structure preserved
    text = soup.get_text(separator="\n")

    # Clean up excessive whitespace
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines)


def read_file(file) -> tuple[str, str]:
    """
    Read file based on extension.
    Returns (content, file_type)
    """
    filename = file.name.lower()

    if filename.endswith(".pdf"):
        return read_pdf(file), "PDF"
    elif filename.endswith(".docx"):
        return read_docx(file), "Word"
    elif filename.endswith(".csv"):
        return read_csv(file), "CSV"
    elif filename.endswith(".xlsx"):
        return read_excel(file), "Excel"
    elif filename.endswith(".html") or filename.endswith(".htm"):
        return read_html(file), "HTML"
    else:  # txt, md, or unknown text
        return read_txt(file), "Text"


# ============== Core RAG Functions ==============


def chunk_text(
    text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP
) -> list[str]:
    """Split text into overlapping chunks."""
    words = text.split()
    chunks = []

    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i : i + chunk_size])
        if chunk.strip():
            chunks.append(chunk)

    return chunks


def get_embeddings(texts: list[str], api_key: str) -> np.ndarray:
    """Get embeddings from OpenAI API."""
    client = openai.OpenAI(api_key=api_key)

    # Process in batches of 100 (API limit)
    all_embeddings = []
    batch_size = 100

    for i in range(0, len(texts), batch_size):
        batch = texts[i : i + batch_size]
        response = client.embeddings.create(model=EMBEDDING_MODEL, input=batch)
        embeddings = [item.embedding for item in response.data]
        all_embeddings.extend(embeddings)

    return np.array(all_embeddings, dtype="float32")


def create_faiss_index(embeddings: np.ndarray) -> faiss.IndexFlatL2:
    """Create a FAISS index from embeddings."""
    dimension = embeddings.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings)
    return index


def search_similar(
    query: str,
    index: faiss.IndexFlatL2,
    chunks: list[str],
    api_key: str,
    top_k: int = 3,
) -> list[str]:
    """Find most similar chunks to the query."""
    query_embedding = get_embeddings([query], api_key)
    distances, indices = index.search(query_embedding, top_k)

    results = [chunks[i] for i in indices[0] if i < len(chunks)]
    return results


def generate_answer(query: str, context_chunks: list[str], api_key: str) -> str:
    """Generate answer using OpenAI chat completion with retrieved context."""
    client = openai.OpenAI(api_key=api_key)

    context = "\n\n---\n\n".join(context_chunks)

    system_prompt = """You are a helpful assistant that answers questions based on the provided context.
    
Rules:
- Answer ONLY based on the context provided
- If the context doesn't contain relevant information, say "I couldn't find relevant information in the documents."
- Be concise and direct
- Quote relevant parts when appropriate
- If the question asks for data from tables/spreadsheets, present it clearly"""

    user_prompt = f"""Context:
{context}

---

Question: {query}

Answer:"""

    response = client.chat.completions.create(
        model=CHAT_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.3,
        max_tokens=1000,
    )

    return response.choices[0].message.content


# ============== Streamlit UI ==============


def main():
    st.set_page_config(page_title="RAG Q&A System", page_icon="📚", layout="wide")

    st.title("📚 RAG Document Q&A System")
    st.markdown("*Upload documents, ask questions, get AI-powered answers*")

    # Initialize session state
    if "chunks" not in st.session_state:
        st.session_state.chunks = []
    if "index" not in st.session_state:
        st.session_state.index = None
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "documents_loaded" not in st.session_state:
        st.session_state.documents_loaded = False
    if "doc_stats" not in st.session_state:
        st.session_state.doc_stats = []

    # Sidebar for configuration

    with st.sidebar:
        st.header("⚙️ Configuration")

        # Get API key from secrets
        api_key = get_api_key()

        if api_key:
            st.success("✅ API Key loaded")
        else:
            st.error("❌ API Key not found")
            st.markdown(
                """
            **Setup required:**
            
            Create `.streamlit/secrets.toml`:
            ```toml
            OPENAI_API_KEY = "sk-your-key-here"
            ```
            
            Or set environment variable:
            ```bash
            export OPENAI_API_KEY="sk-your-key"
            ```
            """
            )

        st.divider()

        # Document upload section
        st.header("📄 Add Documents")

        # Show supported formats
        with st.expander("Supported Formats"):
            st.markdown(
                """
            | Format | Extension |
            |--------|-----------|
            | Plain Text | `.txt` |
            | Markdown | `.md` |
            | PDF | `.pdf` |
            | Word | `.docx` |
            | CSV | `.csv` |
            | Excel | `.xlsx` |
            | HTML | `.html` |
            """
            )

        # File upload
        uploaded_files = st.file_uploader(
            "Upload files",
            type=SUPPORTED_TYPES,
            accept_multiple_files=True,
            help="Upload one or more documents",
        )

        # Or paste text
        pasted_text = st.text_area(
            "Or paste text directly", height=120, placeholder="Paste content here..."
        )

        # Advanced settings
        with st.expander("Advanced Settings"):
            chunk_size = st.slider("Chunk size (words)", 100, 1000, CHUNK_SIZE, 50)
            chunk_overlap = st.slider(
                "Chunk overlap (words)", 0, 200, CHUNK_OVERLAP, 10
            )
            top_k = st.slider("Results to retrieve", 1, 10, 3)

        # Process button
        if st.button("🔄 Process Documents", type="primary", disabled=not api_key):
            all_text = ""
            doc_stats = []

            # Process uploaded files
            if uploaded_files:
                progress = st.progress(0, "Processing files...")

                for i, file in enumerate(uploaded_files):
                    try:
                        content, file_type = read_file(file)
                        word_count = len(content.split())
                        all_text += f"\n\n--- Document: {file.name} ---\n\n{content}"
                        doc_stats.append(
                            {"name": file.name, "type": file_type, "words": word_count}
                        )
                        progress.progress(
                            (i + 1) / len(uploaded_files), f"Processed {file.name}"
                        )
                    except Exception as e:
                        st.error(f"Error reading {file.name}: {str(e)}")

                progress.empty()

            # Add pasted text
            if pasted_text.strip():
                all_text += f"\n\n--- Pasted Text ---\n\n{pasted_text}"
                doc_stats.append(
                    {
                        "name": "Pasted text",
                        "type": "Text",
                        "words": len(pasted_text.split()),
                    }
                )

            if all_text.strip():
                with st.spinner("Creating embeddings and index..."):
                    try:
                        # Chunk the text
                        chunks = chunk_text(all_text, chunk_size, chunk_overlap)
                        st.session_state.chunks = chunks

                        # Create embeddings and index
                        embeddings = get_embeddings(chunks, api_key)
                        st.session_state.index = create_faiss_index(embeddings)
                        st.session_state.documents_loaded = True
                        st.session_state.doc_stats = doc_stats
                        st.session_state.top_k = top_k

                        st.success(f"✅ Processed {len(chunks)} chunks!")

                    except Exception as e:
                        st.error(f"Error: {str(e)}")
            else:
                st.warning("Please upload files or paste text")

        # Show loaded documents
        if st.session_state.documents_loaded:
            st.divider()
            st.subheader("📊 Loaded Documents")

            for doc in st.session_state.doc_stats:
                st.markdown(
                    f"**{doc['name']}**  \n{doc['type']} • {doc['words']:,} words"
                )

            st.markdown(f"**Total chunks:** {len(st.session_state.chunks)}")

        # Clear button
        if st.session_state.documents_loaded:
            st.divider()
            if st.button("🗑️ Clear All"):
                st.session_state.chunks = []
                st.session_state.index = None
                st.session_state.chat_history = []
                st.session_state.documents_loaded = False
                st.session_state.doc_stats = []
                st.rerun()

    # Main chat interface

    if not api_key:
        st.warning(
            "⚠️ OpenAI API key not configured. See sidebar for setup instructions."
        )
        return

    if not st.session_state.documents_loaded:
        st.info("👈 Upload documents in the sidebar, then click 'Process Documents'")

        # Show example
        with st.expander("📖 How it works"):
            st.markdown(
                """
            **RAG (Retrieval-Augmented Generation) Process:**
            
            1. **Upload** — Add your documents (PDF, Word, Excel, CSV, Text, HTML)
            2. **Process** — Documents are split into chunks and converted to embeddings
            3. **Index** — Embeddings are stored in a FAISS vector index
            4. **Query** — Your questions are matched against the document chunks
            5. **Generate** — Relevant chunks are sent to GPT to generate accurate answers
            
            **Supported document types:**
            - 📄 PDF documents
            - 📝 Word documents (.docx)
            - 📊 Spreadsheets (CSV, Excel)
            - 🌐 HTML files
            - 📃 Text and Markdown files
            """
            )
        return

    # Display chat history
    for msg in st.session_state.chat_history:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            if "sources" in msg and msg["sources"]:
                with st.expander("📎 Sources"):
                    for i, source in enumerate(msg["sources"], 1):
                        st.markdown(f"**Chunk {i}:**")
                        display_text = (
                            source[:500] + "..." if len(source) > 500 else source
                        )
                        st.markdown(f"> {display_text}")

    # Chat input
    if query := st.chat_input("Ask a question about your documents..."):
        # Add user message
        st.session_state.chat_history.append({"role": "user", "content": query})

        with st.chat_message("user"):
            st.markdown(query)

        # Generate response
        with st.chat_message("assistant"):
            with st.spinner("Searching and generating answer..."):
                try:
                    # Get top_k from session or default
                    top_k = st.session_state.get("top_k", 3)

                    # Search for relevant chunks
                    relevant_chunks = search_similar(
                        query,
                        st.session_state.index,
                        st.session_state.chunks,
                        api_key,
                        top_k=top_k,
                    )

                    # Generate answer
                    answer = generate_answer(query, relevant_chunks, api_key)

                    st.markdown(answer)

                    # Show sources
                    with st.expander("📎 Sources"):
                        for i, source in enumerate(relevant_chunks, 1):
                            st.markdown(f"**Chunk {i}:**")
                            display_text = (
                                source[:500] + "..." if len(source) > 500 else source
                            )
                            st.markdown(f"> {display_text}")

                    # Save to history
                    st.session_state.chat_history.append(
                        {
                            "role": "assistant",
                            "content": answer,
                            "sources": relevant_chunks,
                        }
                    )

                except Exception as e:
                    st.error(f"Error: {str(e)}")


if __name__ == "__main__":
    main()
